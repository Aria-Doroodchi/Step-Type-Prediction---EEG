# -*- coding: utf-8 -*-
"""Build the 3-state (standing/straight/diagonal) supervisor Word report (.docx),
mirroring the HTML report (build_report.py) and the step-type docx format
(scripts/_report_build_docx.py). python-docx; figures embedded from the state
report's figures dir. Full 32-participant cohort numbers.
"""
from pathlib import Path
import re
from PIL import Image
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(r"C:/Users/Ali D/Documents/ML")
DIR = ROOT / "outputs" / "reports" / "state_3class_2026-06-26"
FIGS = DIR / "figures"
OUT = DIR / "EEG_State_Report.docx"
TITLE_RUNNING = "Reading Motor State from Brain Activity"

BLUE = "1B3A6B"; ACCENT = "1F77B4"; GREEN = "2F7D4F"; WARN = "A85A00"; MUTED = "5B6B7B"


# ---------- low-level helpers (from _report_build_docx.py) ----------
def _shd(fill):
    e = OxmlElement('w:shd'); e.set(qn('w:val'), 'clear'); e.set(qn('w:color'), 'auto'); e.set(qn('w:fill'), fill)
    return e

def shade_cell(cell, fill):
    cell._tc.get_or_add_tcPr().append(_shd(fill))

def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr(); m = OxmlElement('w:tcMar')
    for edge, val in [('top', top), ('bottom', bottom), ('start', left), ('end', right)]:
        x = OxmlElement(f'w:{edge}'); x.set(qn('w:w'), str(val)); x.set(qn('w:type'), 'dxa'); m.append(x)
    tcPr.append(m)

def cell_borders(cell, edges):
    tcPr = cell._tc.get_or_add_tcPr(); b = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        color, size = edges.get(edge, ('FFFFFF', 0))
        x = OxmlElement(f'w:{edge}')
        x.set(qn('w:val'), 'single' if size else 'nil'); x.set(qn('w:sz'), str(size))
        x.set(qn('w:space'), '0'); x.set(qn('w:color'), color); b.append(x)
    tcPr.append(b)

def set_table_width(tbl, widths):
    tbl.allow_autofit = False; tblPr = tbl._tbl.tblPr
    w = OxmlElement('w:tblW'); w.set(qn('w:w'), str(sum(widths))); w.set(qn('w:type'), 'dxa'); tblPr.append(w)
    for row in tbl.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Pt(widths[i] / 20)

def rich(p, text, size=None, color=None, italic=False):
    for part in re.split(r'(\*\*.*?\*\*|\*[^*]+?\*)', text):
        if not part:
            continue
        is_bold = part.startswith('**') and part.endswith('**')
        is_ital = (not is_bold) and part.startswith('*') and part.endswith('*') and len(part) > 2
        if is_bold:
            r = p.add_run(part[2:-2]); r.bold = True
        elif is_ital:
            r = p.add_run(part[1:-1]); r.italic = True
        else:
            r = p.add_run(part)
        if italic: r.italic = True
        if size: r.font.size = Pt(size)
        if color: r.font.color.rgb = RGBColor.from_string(color)
    return p


# ---------- block builders ----------
def heading(doc, text, level):
    h = doc.add_heading(level=level); run = h.add_run(text)
    run.font.color.rgb = RGBColor.from_string(BLUE); return h

def para(doc, text, size=11, space_after=8):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(space_after)
    rich(p, text, size=size); return p

def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet'); rich(p, text); return p

def figure(doc, fname, caption, num, width_in=6.5):
    Image.open(FIGS / fname)  # validate present
    pic_p = doc.add_paragraph(); pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.add_run().add_picture(str(FIGS / fname), width=Inches(width_in))
    cap = doc.add_paragraph(); cap.paragraph_format.space_before = Pt(3); cap.paragraph_format.space_after = Pt(14)
    r = cap.add_run(f"Figure {num}.  "); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor.from_string(BLUE)
    rr = cap.add_run(caption); rr.font.size = Pt(9); rr.font.color.rgb = RGBColor.from_string(MUTED)

def callout(doc, label, lines, accent=ACCENT, fill="F4F7FB"):
    tbl = doc.add_table(rows=1, cols=1); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(tbl, [9360]); cell = tbl.cell(0, 0)
    shade_cell(cell, fill); set_cell_margins(cell, 120, 120, 200, 160)
    cell_borders(cell, {'left': (accent, 28), 'top': ('E4E9EF', 4), 'bottom': ('E4E9EF', 4), 'right': ('E4E9EF', 4)})
    p0 = cell.paragraphs[0]; p0.paragraph_format.space_after = Pt(4)
    rl = p0.add_run(label.upper()); rl.bold = True; rl.font.size = Pt(8.5); rl.font.color.rgb = RGBColor.from_string(accent)
    for ln in lines:
        bp = cell.add_paragraph(); bp.paragraph_format.space_after = Pt(2); rich(bp, ln, size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2); return tbl

def metric_cards(doc, cards):
    tbl = doc.add_table(rows=2, cols=len(cards)); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cw = 9360 // len(cards); set_table_width(tbl, [cw] * len(cards))
    for i, (big, lab) in enumerate(cards):
        top = tbl.cell(0, i); bot = tbl.cell(1, i)
        for c in (top, bot):
            shade_cell(c, "FFFFFF"); set_cell_margins(c, 90, 90, 120, 120)
            cell_borders(c, {e: ('DCE3EA', 4) for e in ('top', 'left', 'bottom', 'right')})
        pt = top.paragraphs[0]; pt.alignment = WD_ALIGN_PARAGRAPH.CENTER; pt.paragraph_format.space_after = Pt(0)
        rt = pt.add_run(big); rt.bold = True; rt.font.size = Pt(22); rt.font.color.rgb = RGBColor.from_string(BLUE)
        pb = bot.paragraphs[0]; pb.alignment = WD_ALIGN_PARAGRAPH.CENTER; pb.paragraph_format.space_after = Pt(0)
        rb = pb.add_run(lab); rb.font.size = Pt(8.5); rb.font.color.rgb = RGBColor.from_string(MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def simple_table(doc, headers, rows, widths):
    tbl = doc.add_table(rows=1, cols=len(headers)); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(tbl, widths)
    for i, htext in enumerate(headers):
        c = tbl.cell(0, i); shade_cell(c, "EAF0F7"); set_cell_margins(c)
        cell_borders(c, {e: ('CCD6E0', 4) for e in ('top', 'left', 'bottom', 'right')})
        p = c.paragraphs[0]; r = p.add_run(htext); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor.from_string(BLUE)
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            set_cell_margins(cells[i]); cell_borders(cells[i], {e: ('DCE3EA', 4) for e in ('top', 'left', 'bottom', 'right')})
            rich(cells[i].paragraphs[0], val, size=10)
    return tbl

def add_toc(doc):
    p = doc.add_paragraph(); run = p.add_run()
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = "Open in Word and choose Update Field to build the contents."
    f3 = OxmlElement('w:fldChar'); f3.set(qn('w:fldCharType'), 'end')
    for el in (f1, instr, f2, t, f3): run._r.append(el)

def set_update_fields(doc):
    uf = OxmlElement('w:updateFields'); uf.set(qn('w:val'), 'true'); doc.settings.element.append(uf)

def page_number_footer(section):
    p = section.footer.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(TITLE_RUNNING + "   -   Page ").font.size = Pt(8.5)
    r = p.add_run()
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = 'PAGE'
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end')
    for el in (f1, it, f2): r._r.append(el)
    r.font.size = Pt(8.5)
    for run in p.runs: run.font.color.rgb = RGBColor.from_string(MUTED)


# ---------- document ----------
def build():
    doc = Document()
    normal = doc.styles['Normal']; normal.font.name = 'Calibri'; normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string('1F2933')
    for hid, sz in [('Heading 1', 17), ('Heading 2', 14)]:
        st = doc.styles[hid]; st.font.name = 'Calibri'; st.font.size = Pt(sz); st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(BLUE)
    sec = doc.sections[0]
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    sec.top_margin = sec.bottom_margin = Inches(1); sec.left_margin = sec.right_margin = Inches(1)
    page_number_footer(sec)

    # ----- TITLE -----
    eb = doc.add_paragraph(); eb.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = eb.add_run("EEG  -  MOTOR STATE  -  MACHINE LEARNING")
    r.font.size = Pt(10); r.font.color.rgb = RGBColor.from_string(ACCENT); r.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Reading Motor State"); r.bold = True; r.font.size = Pt(30); r.font.color.rgb = RGBColor.from_string(BLUE)
    r.add_break(); r2 = title.add_run("from Brain Activity"); r2.bold = True; r2.font.size = Pt(30); r2.font.color.rgb = RGBColor.from_string(BLUE)
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Single-trial classification of three motor states - quiet stance, straight stepping, "
                     "and diagonal stepping - from 64-channel scalp EEG, with a per-participant model and "
                     "nested cross-validation. A companion to the straight-vs-diagonal step-type project.")
    rs.font.size = Pt(13); rs.font.color.rgb = RGBColor.from_string(MUTED); rs.italic = True
    for _ in range(2): doc.add_paragraph()
    for lab, val in [("Prepared by:", "[your name]"), ("For:", "[supervisor name]"),
                     ("Affiliation:", "Sunnybrook Research Institute"), ("Date:", "June 2026 - full 32-participant cohort")]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(3)
        rl = p.add_run(lab + "  "); rl.bold = True; rl.font.size = Pt(11)
        rv = p.add_run(val); rv.font.size = Pt(11)
        if val.startswith("["): rv.font.highlight_color = WD_COLOR_INDEX.YELLOW
    doc.add_page_break()

    heading(doc, "What's in this report", 1); add_toc(doc); doc.add_page_break()

    # ----- 1 -----
    heading(doc, "1   The short version", 1)
    para(doc, "This is the three-way sibling of the step-type project. Instead of asking *straight or "
         "diagonal?*, it poses a broader decoding question: **from a single two-second epoch of 64-channel "
         "EEG, can we recover the participant's motor state - quiet stance, straight stepping, or diagonal "
         "stepping?** Each participant is modelled separately and scored only on their own held-out epochs.")
    metric_cards(doc, [("0.88", "average score (macro-AUC), where 0.50 = chance"),
                       ("73%", "three-way accuracy (chance = 33%)"),
                       ("32", "participants (full cohort), each with their own model")])
    callout(doc, "Headline",
            ["Across all thirty-two participants, the model separates the three motor states with an average "
             "**macro-AUC of 0.878** (chance = 0.50) and **73% accuracy** (chance = 33%), and it does so "
             "honestly - the gap between its optimistic internal score and its real held-out score is tiny "
             "(0.05). **Standing is almost always identified correctly; telling a straight step from a "
             "diagonal one is the hard part.** A controlled comparison (Section 7) shows a planned "
             "*foot-stimulation* feature adds essentially nothing on top of the basic scalp features, while "
             "the *brain-source* feature adds a small but real boost to the hard straight-vs-diagonal "
             "distinction."], accent=GREEN, fill="F1F8F3")

    # ----- 2 -----
    heading(doc, "2   What the project is trying to do", 1)
    para(doc, "The step-type project showed that the pre-movement **contingent negative variation (CNV)** - "
         "the slow anticipatory potential of the warning-imperative foreperiod - carries a faint trace of "
         "*which way* a participant is about to step. This project widens the scope from two step-types to "
         "**three motor states**: quiet stance, straight stepping, and diagonal stepping, decoded from the "
         "ongoing sensorimotor EEG rather than only the preparatory window. It is a step toward decoding a "
         "richer state repertoire from a single montage.")
    para(doc, "With **three** classes, a naive model would reach only **one third (33%)** accuracy; anything "
         "reliably above that indicates the EEG genuinely separates the states. We retain the project's core "
         "design: **one model per participant**, each evaluated only on epochs it never saw during training "
         "or tuning.")

    # ----- 3 -----
    heading(doc, "3   The three states and where the data come from", 1)
    para(doc, "Each participant contributes two continuous 64-channel recordings (BioSemi, 1024 or 2048 Hz) "
         "that between them supply the three states. Throughout, a single **2 s epoch** of the 64-channel "
         "montage is one exemplar for the classifier.")
    figure(doc, "fig_paradigm.png", "The three-state paradigm and epoching. Stepping (top, Stim recording): "
           "a warning-imperative structure in which an S1 direction cue (event code 256 = straight, 512 = "
           "diagonal) precedes the response / gait onset (code 96); the analysis epoch is response-locked "
           "over [-0.1, +2.0] s. Standing (bottom, Standing recording): no cue structure, so non-overlapping "
           "2 s windows are tiled from the continuous record and count-balanced to the stepping trials. A "
           "continuous foot-sole electrical-stimulation train (gold ticks; ~2 Hz, 0.52 s ISI) runs through "
           "both recordings; each pulse evokes a low-amplitude vertex somatosensory evoked potential (inset; "
           "P50/N90), with the stimulator trigger offset from the true cutaneous pulse by a calibrated "
           "273 ms.", 1)
    bullet(doc, "**Straight / diagonal stepping** - from the *Stim* recording. The S1 cue signals the "
           "direction; the participant steps ~2 s later. Epochs are locked to the paired response (code 96) "
           "over [-0.1, +2.0] s, exactly as in the step-type project (~40 trials per direction).")
    bullet(doc, "**Quiet stance** - from the *Standing* recording. With no cue structure, non-overlapping "
           "2 s windows are tiled from the continuous record and down-sampled to the stepping count, giving "
           "a balanced ~40/40/40 three-class problem per participant.")
    para(doc, "A methodological feature of this dataset is that a foot-sole electrical stimulator runs "
         "*continuously* - at ~2 Hz (0.52 s ISI) - across *both* the stepping and standing recordings. "
         "Because the stimulation train is present and matched in every state, it is not a between-state cue "
         "in itself; rather, each pulse evokes a small cortical **somatosensory evoked potential (SEP)** at "
         "the vertex, which we exploited as an additional per-epoch feature family (Section 5). The "
         "stimulator trigger is offset from the true cutaneous pulse by a fixed **273 ms** (calibrated "
         "cohort-wide; 95% CI [271.6, 276.1] ms), applied per recording as round(0.273 x sfreq) samples to "
         "accommodate the mixed 1024/2048 Hz sampling.")

    # ----- 4 -----
    heading(doc, "4   The prediction engine: a 3-way XGBoost model", 1)
    para(doc, "As in the step-type project, the workhorse is **XGBoost** - a model built from many small "
         "**decision trees** (yes/no flowcharts over the measured numbers) added together, each new tree "
         "correcting the previous ones' mistakes. The only change for this project is the output: instead of "
         "one probability (“diagonal?”), the model now produces **three probabilities that add to "
         "100%** - one per state - and predicts whichever is largest.")
    figure(doc, "fig_xgb_tree.png", "A single decision tree - the building block. Each box asks a yes/no "
           "question about one measurement; following the answers leads to a verdict. The model invents its "
           "own thresholds from the continuous data.", 2, width_in=5.6)
    figure(doc, "fig_xgb_boosting.png", "Gradient boosting: hundreds of simple trees are added together, "
           "each correcting the running total left by the trees before it. For this project the combined "
           "scores are turned into three competing probabilities (standing / straight / diagonal).", 3)
    para(doc, "Performance is summarised with **macro-AUC** - the natural three-class version of the AUC "
         "score. It asks, for each state in turn, “how well does the model separate this state from the "
         "other two?”, then averages. As before, **0.50 is chance and 1.00 is perfect.**")

    # ----- 5 -----
    heading(doc, "5   What the model looks at: the features", 1)
    para(doc, "Each 2 s epoch is reduced to a large feature vector of interpretable single-trial measures. "
         "We reuse the four families from the step-type pipeline, each computed per channel and per "
         "**62.5 ms (1/16 s) time bin**:")
    for t in ["**Amplitude** - the mean potential within each bin: the binned single-trial ERP.",
              "**Slopes** - the within-bin temporal gradient (first derivative), indexing the *rate* of "
              "change rather than the absolute level.",
              "**Spectral power** - single-trial band power from a Morlet-wavelet time-frequency "
              "decomposition, summarised in the canonical bands (delta, theta, alpha, beta, gamma).",
              "**Source estimates** - a distributed inverse solution (eLORETA) projected onto the fsaverage "
              "template cortex and parcellated with the Destrieux atlas (aparc.a2009s, ~150 ROIs)."]:
        bullet(doc, t)
    figure(doc, "fig_binning.png", "Turning a continuous waveform into features. Each 2 s epoch is divided "
           "into non-overlapping 62.5 ms bins and summarised per bin; across channels, bins, bands and "
           "source ROIs this yields ~19,500 candidate features per epoch.", 4)
    callout(doc, "A note on referencing - CSD (surface Laplacian)",
            ["The scalp-electrode features (amplitude, slopes, power) are computed not on the raw "
             "average-referenced potentials but on their **current source density (CSD)** transform - the "
             "surface Laplacian of the scalp field. CSD is a reference-free spatial high-pass filter that "
             "sharpens the topography of superficial radial generators and suppresses volume-conducted and "
             "far-field activity; it is expressed here in arbitrary units. The trade-off is that, as a "
             "spatial second derivative, CSD also amplifies high-spatial-frequency and broadband noise - "
             "which is why the grand-average display below (Figure 5) is drawn from the average-reference "
             "(uV) epochs rather than the CSD ones, even though the classifier's window features use CSD."])
    para(doc, "**The novel block - a per-epoch foot-SEP.** Because the sole is stimulated throughout, we "
         "added a fifth feature family unique to this project, read from the *average-reference* (non-CSD) "
         "epochs - the vertex foot-SEP is a deep, largely tangential paracentral generator that the surface "
         "Laplacian would attenuate. For each analysis epoch we average the SEPs of its own in-epoch pulses "
         "(~4 per 2 s window), baseline-correct (-50 to 0 ms), and read fixed a-priori vertex components - "
         "an early positivity (**P50**, 40-50 ms window) and the dominant negativity (**N90**, 75-90 ms "
         "window; grand-average trough ~65-75 ms), plus peak-to-peak and RMS over 15-130 ms, across the "
         "vertex montage [Cz, C1, C2, FCz, CPz], with the read window starting >=15 ms to exclude the "
         "stimulus artifact. (*P50/N90* are latency-defined labels for the vertex foot-SEP peaks, not the "
         "canonical tibial-nerve P37/N45 nomenclature - plantar cutaneous stimulation evokes a later, sub-uV "
         "complex.) The block is **leakage-safe by construction**: every epoch's SEP is built only from that "
         "epoch's own pulses, never a per-condition average broadcast across trials.")
    figure(doc, "fig_condition_erp.png", "Grand-average vertex window ERP per state (average reference, uV; "
           "zero-phase 6 Hz low-pass for display), across all thirty-two participants. Standing (blue) stays "
           "near baseline, whereas both stepping states show a large early positivity (~0.2 s, around "
           "response / gait onset) followed by a sustained negativity that is deepest for diagonal. The two "
           "stepping waveforms share a similar early positivity but diverge through the sustained negativity; "
           "at the single-trial level, though, quiet stance is trivially separable whereas straight vs "
           "diagonal remains subtle. The classifier's window features are the CSD transform of these "
           "signals; the display is in uV for legibility.", 5)
    figure(doc, "fig_sep.png", "Grand-average vertex foot-SEP per state (average reference, uV), time-locked "
           "to the true (offset-corrected) pulse. The evoked complex is low-amplitude (sub-uV): a weak P50 "
           "positivity and a dominant N90 trough. The three states are near-identical - consistent with the "
           "ablation (Section 7), where the SEP block adds nothing beyond the window features.", 6)

    # ----- 6 -----
    heading(doc, "6   How the model is built and tested", 1)
    para(doc, "The testing machinery is identical to the step-type project, because honesty matters more "
         "than any single number. Two principles run through it: **never let the model see the data it will "
         "be scored on**, and **trim the ~19,500 raw features down to the few dozen that matter**, doing "
         "both *inside* each training fold so nothing leaks.")
    figure(doc, "fig_nestedcv.png", "Nested cross-validation. An outer split sets aside a test block used "
           "only for the final score; all choices are made on a separate inner split. Because the test block "
           "is never used to make a choice, the score is not inflated.", 7)
    figure(doc, "fig_pipeline.png", "The feature funnel, applied fresh inside every training fold: from "
           "~19,500 candidates down to a few dozen, via a correlation filter, a quick statistical screen, "
           "stability selection, and gain pruning. The three-class run uses a multiclass-safe version of "
           "each step.", 8)
    para(doc, "To keep the three states evenly matched, the standing windows are down-sampled to the number "
         "of stepping trials, and the cross-validation is **stratified** so every split contains all three "
         "states. Each per-person model is then scored on held-out epochs it never saw.")

    # ----- 7 -----
    heading(doc, "7   Results: which states are separable, and what helps", 1)
    para(doc, "The headline numbers come from all thirty-two participants with the full honest test. To find "
         "out *which features actually matter*, we ran the same analysis four times, each with a different "
         "feature set - an **ablation**. This is the most informative result in the report.")
    figure(doc, "fig_ablation.png", "The ablation. Each bar is the cohort macro-AUC for a different feature "
           "set (chance = 0.50, dashed). The full set (combined) and the set without the foot-SEP (window) "
           "are identical; dropping the eLORETA source features too (electrode) costs a small but real "
           "~0.015 macro-AUC. The foot-SEP on its own (sep) lands only modestly above chance.", 9)
    simple_table(doc, ["Feature set", "What's included", "Macro-AUC", "3-way accuracy"], [
        ["**Combined**", "everything (scalp + brain-source + foot-SEP)", "**0.878**", "73%"],
        ["**Window**", "scalp + brain-source (no foot-SEP)", "0.877", "73%"],
        ["**Electrode**", "scalp features only (amplitude, slopes, power)", "0.862", "71%"],
        ["**SEP only**", "just the foot-stimulation feature", "0.584", "40%"],
    ], [1700, 4360, 1650, 1650])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    callout(doc, "Two clean answers",
            ["**Does the new foot-SEP feature add anything?** No. Combined and Window score identically "
             "(0.878 vs 0.877), so once the ordinary scalp features are present, the foot-SEP is redundant - "
             "although on its own it does carry a weak signal (0.58). **Do the expensive brain-source "
             "features earn their keep?** Yes, modestly - removing them drops the score by 0.015 (0.877 to "
             "0.862), and more tellingly it costs ~4 points of accuracy on the hard straight-vs-diagonal "
             "distinction (63/61% to 59/57%). In an early 8-person preview this gap looked negligible; "
             "across the full cohort the brain-source contribution is small but consistent."],
            accent=GREEN, fill="F1F8F3")
    para(doc, "The **confusion matrix** shows *where* the accuracy comes from: rows are the true state, "
         "columns are the model's guess, and the diagonal counts correct calls.")
    figure(doc, "fig_confusion.png", "Three-by-three confusion matrix across all thirty-two participants. "
           "Standing is identified almost perfectly (~96% correct). The errors concentrate in the "
           "bottom-right block: straight and diagonal steps are most often confused with each other "
           "(~61-63% correct each), not with standing.", 10, width_in=5.0)
    para(doc, "So the three states are **not** equally easy. Per-class accuracy is about **96% for standing, "
         "63% for straight, and 61% for diagonal**. Standing is trivially separable from movement; the "
         "genuinely hard sub-problem - telling a straight step from a diagonal one - is exactly the original "
         "step-type question, and it sits modestly above its own two-way chance line.")
    figure(doc, "fig_perparticipant.png", "Each participant's own model. Top: three-class macro-AUC (chance "
           "0.50, dashed). Bottom: overall accuracy (chance 0.333, dashed). Every participant is clearly "
           "above chance on both measures, with the usual person-to-person variation.", 11)
    callout(doc, "The honesty check",
            ["An overfit model scores far higher on its own tuning data than on unseen data. Here the "
             "inner-vs-outer gap is only **0.05**, and every participant clears chance on the truly held-out "
             "epochs - so the 0.878 figure is a fair, not inflated, estimate for the full 32-participant "
             "cohort."],
            accent=GREEN, fill="F1F8F3")

    # ----- 8 -----
    heading(doc, "8   An honest caveat: the standing-stepping movement asymmetry", 1)
    callout(doc, "Read this before over-interpreting standing",
            ["Near-perfect detection of standing should not be read as near-perfect decoding of a *cortical* "
             "state. Because the foot-stimulation train is continuous and matched across all three states "
             "(Section 3), it is *not* a between-state confound here - and we additionally blank the "
             "+/-12 ms stimulus-artifact intervals before extracting the window features. What still "
             "separates standing from stepping is the obvious factor: **gross movement**. Stepping entails "
             "whole-body motion, and with it postural EMG, motion and cable artifact, and genuine "
             "sensorimotor cortical engagement; quiet stance has none of these. Some of the standing-stepping "
             "separability is therefore **non-neural** (movement-related artifact) rather than a pure "
             "motor-state readout - an asymmetry no reference scheme fully removes.",
             "This is precisely why the **straight-vs-diagonal** contrast is the scientifically clean one: "
             "both are stepping, with matched movement *and* matched stimulation, so any separation there "
             "reflects genuine direction-related sensorimotor activity - the original step-type signal, now "
             "embedded in the three-class setting. Partitioning the standing result into movement-related "
             "artifact versus true motor-state signal (e.g. via EMG / accelerometer regressors or "
             "artifact-matched controls) is the natural next check."],
            accent=WARN, fill="FDF6EE")

    # ----- 9 -----
    heading(doc, "9   Summary and next steps", 1)
    para(doc, "In summary: **a single 2 s EEG epoch separates quiet stance, straight stepping and diagonal "
         "stepping well above chance** (macro-AUC 0.878, accuracy 73% vs 33% chance), under an honest nested "
         "cross-validation estimate. Standing is near-ceiling - subject to the movement-asymmetry caveat "
         "above - while the meaningful, artifact-matched problem is straight vs diagonal. The ablation was "
         "decisive: the per-epoch foot-SEP adds nothing beyond the window features, whereas the eLORETA "
         "source estimates give a small but consistent lift (+0.015 macro-AUC) to the hard "
         "straight-vs-diagonal distinction.")
    para(doc, "**Where this goes next:**")
    bullet(doc, "**The full 32-participant cohort is complete.** Ablation-guided production feature set: "
           "**keep the eLORETA source block** (a small, consistent lift to straight-vs-diagonal) and "
           "**drop the foot-SEP block** (redundant against the window features).")
    bullet(doc, "**Partition the standing result** into movement-related artifact versus genuine motor-state "
           "signal - e.g. with EMG / accelerometer regressors or artifact-matched controls - since the "
           "continuous, matched stimulation already rules out a stimulation-rhythm explanation.")
    bullet(doc, "**Focus on straight-vs-diagonal**, the artifact-matched core problem, and relate it back to "
           "the CNV step-type findings.")

    # ----- GLOSSARY -----
    heading(doc, "Glossary (modelling & statistics)", 1)
    for term, defn in [
        ("Feature", "A single number summarising one aspect of an epoch (e.g. band power at one channel in "
         "one time bin)."),
        ("Macro-AUC", "Macro-averaged one-vs-rest area under the ROC curve: for each state, how well it "
         "separates from the other two, then averaged. 0.50 = chance, 1.00 = perfect."),
        ("Accuracy", "Fraction of epochs assigned the correct state. Chance is 1/3 (33%) for three balanced "
         "classes."),
        ("Ablation", "Re-running the analysis with feature blocks removed, to isolate each block's "
         "contribution."),
        ("Confusion matrix", "A table of true versus predicted state; the diagonal counts correct calls, the "
         "off-diagonal the errors."),
        ("Overfitting", "When a model fits its training data rather than a generalisable rule - high "
         "internal scores, poor held-out scores."),
        ("Nested cross-validation", "An outer split scores the finished model; a separate inner split makes "
         "all tuning and feature-selection choices, so the reported score is not inflated."),
        ("XGBoost", "The classifier used here: an additive ensemble of gradient-boosted decision trees."),
    ]:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
        rt = p.add_run(term + " - "); rt.bold = True; rt.font.color.rgb = RGBColor.from_string('16324F')
        p.add_run(defn)

    set_update_fields(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print("DOCX written:", OUT, "(%.0f KB)" % (OUT.stat().st_size / 1024))


if __name__ == "__main__":
    build()
