# -*- coding: utf-8 -*-
"""Build the supervisor Word report (.docx) with python-docx, mirroring the HTML."""
import re
from pathlib import Path
from PIL import Image
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(r"C:/Users/Ali D/Documents/ML")
DIR = ROOT / "outputs" / "reports" / "supervisor_2026-06-25"
FIGS = DIR / "figures"
OUT = DIR / "EEG_StepType_Report.docx"

BLUE = "1B3A6B"; ACCENT = "1F77B4"; GREEN = "2F7D4F"; WARN = "A85A00"; MUTED = "5B6B7B"
CONTENT_W = Inches(6.5)

# ---------- low-level helpers ----------
def _shd(fill):
    e = OxmlElement('w:shd'); e.set(qn('w:val'), 'clear'); e.set(qn('w:color'), 'auto'); e.set(qn('w:fill'), fill)
    return e

def shade_cell(cell, fill):
    cell._tc.get_or_add_tcPr().append(_shd(fill))

def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement('w:tcMar')
    for edge, val in [('top', top), ('bottom', bottom), ('start', left), ('end', right)]:
        x = OxmlElement(f'w:{edge}'); x.set(qn('w:w'), str(val)); x.set(qn('w:type'), 'dxa'); m.append(x)
    tcPr.append(m)

def cell_borders(cell, edges):
    """edges: dict edge->(color,size) ; size in eighths of a point."""
    tcPr = cell._tc.get_or_add_tcPr()
    b = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        color, size = edges.get(edge, ('FFFFFF', 0))
        x = OxmlElement(f'w:{edge}')
        x.set(qn('w:val'), 'single' if size else 'nil'); x.set(qn('w:sz'), str(size))
        x.set(qn('w:space'), '0'); x.set(qn('w:color'), color)
        b.append(x)
    tcPr.append(b)

def set_table_width(tbl, widths):
    tbl.allow_autofit = False
    tblPr = tbl._tbl.tblPr
    w = OxmlElement('w:tblW'); w.set(qn('w:w'), str(sum(widths))); w.set(qn('w:type'), 'dxa'); tblPr.append(w)
    for row in tbl.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Pt(widths[i] / 20)

def rich(p, text, size=None, color=None, italic=False):
    """Render text with **bold** and *italic* spans."""
    for part in re.split(r'(\*\*.*?\*\*|\*[^*]+?\*)', text):
        if not part:
            continue
        is_bold = part.startswith('**') and part.endswith('**')
        is_ital = (not is_bold) and part.startswith('*') and part.endswith('*') and len(part) > 2
        if is_bold:
            r = p.add_run(part[2:-2]); r.bold = True
        elif is_ital:
            r = p.add_run(part[1:-1]); r.italic = True
        else:
            r = p.add_run(part)
        if italic:
            r.italic = True
        if size: r.font.size = Pt(size)
        if color: r.font.color.rgb = RGBColor.from_string(color)
    return p

# ---------- block builders ----------
def heading(doc, text, level):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.color.rgb = RGBColor.from_string(BLUE)
    return h

def para(doc, text, size=11, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    rich(p, text, size=size)
    return p

def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    rich(p, text)
    return p

def figure(doc, fname, caption, num, width_in=6.5):
    w, h = Image.open(FIGS / fname).size
    pic_p = doc.add_paragraph(); pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.add_run().add_picture(str(FIGS / fname), width=Inches(width_in))
    cap = doc.add_paragraph(); cap.paragraph_format.space_before = Pt(3); cap.paragraph_format.space_after = Pt(14)
    r = cap.add_run(f"Figure {num}.  "); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor.from_string(BLUE)
    rr = cap.add_run(caption); rr.font.size = Pt(9); rr.font.color.rgb = RGBColor.from_string(MUTED)

def callout(doc, label, lines, accent=ACCENT, fill="F4F7FB"):
    tbl = doc.add_table(rows=1, cols=1); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(tbl, [9360])
    cell = tbl.cell(0, 0)
    shade_cell(cell, fill); set_cell_margins(cell, 120, 120, 200, 160)
    cell_borders(cell, {'left': (accent, 28), 'top': ('E4E9EF', 4), 'bottom': ('E4E9EF', 4), 'right': ('E4E9EF', 4)})
    p0 = cell.paragraphs[0]; p0.paragraph_format.space_after = Pt(4)
    rl = p0.add_run(label.upper()); rl.bold = True; rl.font.size = Pt(8.5); rl.font.color.rgb = RGBColor.from_string(accent)
    for i, ln in enumerate(lines):
        bp = cell.add_paragraph(); bp.paragraph_format.space_after = Pt(2)
        rich(bp, ln, size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return tbl

def metric_cards(doc, cards):
    tbl = doc.add_table(rows=2, cols=len(cards)); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cw = 9360 // len(cards); set_table_width(tbl, [cw] * len(cards))
    for i, (big, lab) in enumerate(cards):
        top = tbl.cell(0, i); bot = tbl.cell(1, i)
        for c in (top, bot):
            shade_cell(c, "FFFFFF"); set_cell_margins(c, 90, 90, 120, 120)
            cell_borders(c, {e: ('DCE3EA', 4) for e in ('top', 'left', 'bottom', 'right')})
        pt = top.paragraphs[0]; pt.alignment = WD_ALIGN_PARAGRAPH.CENTER; pt.paragraph_format.space_after = Pt(0)
        rt = pt.add_run(big); rt.bold = True; rt.font.size = Pt(22); rt.font.color.rgb = RGBColor.from_string(BLUE)
        pb = bot.paragraphs[0]; pb.alignment = WD_ALIGN_PARAGRAPH.CENTER; pb.paragraph_format.space_after = Pt(0)
        rb = pb.add_run(lab); rb.font.size = Pt(8.5); rb.font.color.rgb = RGBColor.from_string(MUTED)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def simple_table(doc, headers, rows, widths):
    tbl = doc.add_table(rows=1, cols=len(headers)); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(tbl, widths)
    for i, htext in enumerate(headers):
        c = tbl.cell(0, i); shade_cell(c, "EAF0F7"); set_cell_margins(c)
        cell_borders(c, {e: ('CCD6E0', 4) for e in ('top', 'left', 'bottom', 'right')})
        p = c.paragraphs[0]; r = p.add_run(htext); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor.from_string(BLUE)
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            set_cell_margins(cells[i]); cell_borders(cells[i], {e: ('DCE3EA', 4) for e in ('top', 'left', 'bottom', 'right')})
            p = cells[i].paragraphs[0]; rich(p, val, size=10)
    return tbl

def add_toc(doc):
    p = doc.add_paragraph(); run = p.add_run()
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = "Open in Word and choose “Update Field” to build the contents."
    f3 = OxmlElement('w:fldChar'); f3.set(qn('w:fldCharType'), 'end')
    for el in (f1, instr, f2, t, f3):
        run._r.append(el)

def set_update_fields(doc):
    s = doc.settings.element
    uf = OxmlElement('w:updateFields'); uf.set(qn('w:val'), 'true'); s.append(uf)

def page_number_footer(section):
    p = section.footer.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Reading the Next Step from Brain Activity   ·   Page ").font.size = Pt(8.5)
    r = p.add_run()
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = 'PAGE'
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end')
    for el in (f1, it, f2): r._r.append(el)
    r.font.size = Pt(8.5)
    for run in p.runs: run.font.color.rgb = RGBColor.from_string(MUTED)

# ---------- document ----------
doc = Document()
normal = doc.styles['Normal']; normal.font.name = 'Calibri'; normal.font.size = Pt(11); normal.font.color.rgb = RGBColor.from_string('1F2933')
for hid, sz in [('Heading 1', 17), ('Heading 2', 14)]:
    st = doc.styles[hid]; st.font.name = 'Calibri'; st.font.size = Pt(sz); st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(BLUE)

sec = doc.sections[0]
sec.page_width = Inches(8.5); sec.page_height = Inches(11)
sec.top_margin = sec.bottom_margin = Inches(1); sec.left_margin = sec.right_margin = Inches(1)
page_number_footer(sec)

# ----- TITLE PAGE -----
eyebrow = doc.add_paragraph(); eyebrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
re_ = eyebrow.add_run("EEG  ·  MOVEMENT INTENTION  ·  MACHINE LEARNING")
re_.font.size = Pt(10); re_.font.color.rgb = RGBColor.from_string(ACCENT); re_.bold = True
doc.add_paragraph().paragraph_format.space_after = Pt(0)
title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
rt = title.add_run("Reading the Next Step"); rt.bold = True; rt.font.size = Pt(30); rt.font.color.rgb = RGBColor.from_string(BLUE)
rt.add_break()
rt2 = title.add_run("from Brain Activity"); rt2.bold = True; rt2.font.size = Pt(30); rt2.font.color.rgb = RGBColor.from_string(BLUE)
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run("Predicting whether a person is about to step straight or diagonally — "
                 "from the brain’s “getting ready” signal, before they move.")
rs.font.size = Pt(13); rs.font.color.rgb = RGBColor.from_string(MUTED); rs.italic = True
for _ in range(2): doc.add_paragraph()
meta = [("Prepared by:", "[your name]"), ("For:", "[supervisor name]"),
        ("Affiliation:", "Sunnybrook Research Institute"), ("Date:", "June 2026")]
for lab, val in meta:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(3)
    rl = p.add_run(lab + "  "); rl.bold = True; rl.font.size = Pt(11)
    rv = p.add_run(val); rv.font.size = Pt(11)
    if val.startswith("["): rv.font.highlight_color = WD_COLOR_INDEX.YELLOW
doc.add_page_break()

# ----- TOC -----
heading(doc, "What’s in this report", 1)
add_toc(doc)
doc.add_page_break()

# ----- 1 SHORT VERSION -----
heading(doc, "1   The short version", 1)
para(doc, "This project asks a simple question: **can we tell, from brain activity alone, which way a "
     "person is about to step** — straight ahead or off to the side — *before the movement actually "
     "happens?* We record the electrical activity of the brain (an EEG) during the brief “get ready” "
     "window after a person is told which way to step but before they are told to go, and we train a "
     "computer model to read that signal.")
metric_cards(doc, [("0.71", "average accuracy score (AUC), where 0.50 = a coin flip"),
                   ("30", "participants, each with their own model"),
                   ("19 / 30", "participants clearly above chance (AUC ≥ 0.70)")])
callout(doc, "Headline",
        ["On the most recent full-group analysis, the model separates “straight” from “diagonal” steps "
         "with an average score of **0.71** (a coin-flip would score 0.50, a perfect model 1.00). Results "
         "vary a lot from person to person — some are read almost perfectly, a few barely above chance. "
         "Crucially, the score holds up under a strict honesty check, so it is **not an artefact of the "
         "model “memorising” the data.**"], accent=GREEN, fill="F1F8F3")

# ----- 2 GOAL -----
heading(doc, "2   What the project is trying to do", 1)
para(doc, "When a person prepares a planned movement, a slow negative-going cortical potential develops "
     "over the scalp — the **Contingent Negative Variation (CNV)**, first described by Walter and "
     "colleagues in 1964. It arises during the *foreperiod* of a warning–imperative (S1–S2) paradigm: a "
     "**direction cue** (the warning stimulus, S1) signals which way the participant will step, and "
     "~2 s later a **“go” cue** (the imperative stimulus, S2) triggers the movement. The CNV unfolds in "
     "the S1–S2 interval and is generally taken to index anticipatory attention and motor preparation, "
     "its later phase reflecting sensorimotor readiness over fronto-central cortex.")
figure(doc, "fig_timeline.png", "The structure of a single trial. The CNV develops across the foreperiod "
       "between the direction cue (S1) and the “go” cue (S2). The model only ever uses this pre-movement "
       "interval — it predicts before the imperative stimulus and any overt movement.", 1)
para(doc, "Our goal is to look only at that preparation window and predict the **step type**: a "
     "**straight** step (labelled “One” in the data) versus a **diagonal** step (labelled “Two”). "
     "Because there are exactly two possible answers, this is a **binary classification** problem, and a "
     "model that guessed randomly would be right about half the time. Anything reliably above 50% means "
     "the preparation signal genuinely carries information about the planned direction.")
para(doc, "The data come from **30 participants**, each contributing roughly **80 trials** (about 40 "
     "straight, 40 diagonal), recorded on a **64-channel EEG montage**. The continuous data were "
     "preprocessed with a standard artifact-handling pipeline (line-noise removal, ICA-based artifact "
     "rejection, automated bad-trial repair) and re-expressed as **current-source density** (CSD / "
     "surface Laplacian), which sharpens spatial topography and removes the reference. The record is "
     "segmented into **epochs**: ~2 s single-trial windows time-locked to the direction cue (S1).")

# ----- 3 XGB -----
heading(doc, "3   The prediction engine: what an “XGBoost” model is", 1)
para(doc, "The workhorse of this project is a model called **XGBoost**. To understand it, start with its "
     "building block: the **decision tree**. A decision tree is just a flowchart of yes/no questions. "
     "Each question looks at one number and compares it to a threshold, and the answers guide you down "
     "the tree to a final verdict.")
figure(doc, "fig_xgb_tree.png", "A single decision tree. Each box asks a yes/no question about one "
       "measurement from the brain signal (for example, “was the signal at the top of the head below "
       "this level at 1.2 seconds?”). Following the answers leads to a leaf, which gives a probability "
       "that the step was diagonal.", 2)
heading(doc, "How it handles “continuous” measurements (not just categories)", 2)
para(doc, "A common question is how a method built on yes/no questions can work with our data, which are "
     "not neat categories like “red / blue / green” but **continuous numbers** — voltages, slopes and "
     "powers that can take any value. The answer is that the tree **invents its own thresholds**. During "
     "training it searches through the data and finds the most informative cut-points automatically — "
     "for instance, “signal below −1.8” versus “above −1.8.” The modeller never has to define "
     "categories; the model discovers where the meaningful boundaries lie in the raw numbers. This is "
     "exactly why it suits brain-signal data, which is entirely numeric.")
heading(doc, "From one tree to many: “gradient boosting”", 2)
para(doc, "A single small tree is a weak guesser. XGBoost’s trick — the “boosting” in the name — is to "
     "build **hundreds of small trees in sequence, where each new tree focuses on fixing the mistakes "
     "the previous trees made.** The first tree makes a rough guess; the second learns where the first "
     "went wrong and nudges those cases; the third corrects what remains; and so on. The final "
     "prediction adds up the contributions of all the trees and converts the total into a probability "
     "between 0% and 100%.")
figure(doc, "fig_xgb_boosting.png", "Gradient boosting in a nutshell. Many simple trees are added "
       "together, each one correcting the running total left by the trees before it. The combined score "
       "is squashed into a probability of “diagonal.” In our models there are up to a thousand such "
       "trees.", 3)
para(doc, "This combination — many simple trees, each improving on the last — is what makes XGBoost both "
     "accurate and good at spotting subtle *combinations* of signals (e.g. “high power in one place and "
     "a rising slope in another”). It is one of the most successful methods for the kind of table-shaped "
     "data we have here, where each trial is a row and each measurement is a column.")

# ----- 4 FEATURES -----
heading(doc, "4   What the model looks at: the features", 1)
para(doc, "The model never sees the raw squiggly EEG trace. Instead, we summarise each trial into a long "
     "list of numbers called **features** — each one capturing a specific, interpretable aspect of the "
     "brain signal. We compute four families of features.")
for t in [
    "**Amplitude** — the mean CSD-transformed potential within each time bin at each electrode: the "
    "binned single-trial ERP, i.e. the CNV waveform itself.",
    "**Slopes** — the local temporal gradient (linear slope / first derivative) of the potential within "
    "each bin, sensitive to the *rate* of the CNV’s evolving negativity rather than its absolute level.",
    "**Spectral power (PSD)** — single-trial power from a **Morlet-wavelet** time-frequency "
    "decomposition, summarised in the canonical bands: delta (0.5–4 Hz), theta (4–8), alpha (8–13), "
    "beta (13–30) and gamma (30–40).",
    "**Source localisation** — a distributed inverse solution (**eLORETA**) projected onto the "
    "*fsaverage* template cortex and parcellated with the Destrieux atlas (*aparc.a2009s*, ~150 ROIs), "
    "estimating the cortical generators of the scalp signal.",
]:
    bullet(doc, t)
heading(doc, "The binning architecture: turning a wave into numbers", 2)
para(doc, "To turn each continuous waveform into features, the ~2 s epoch (0–2 s post-S1) is divided into "
     "non-overlapping **62.5 ms (1/16 s) time bins**, and every feature is computed per electrode (or "
     "source ROI), per bin, and — for spectral power — per band. One feature might be “the CSD-transformed "
     "potential at the vertex (Cz) averaged between 0.75 and 0.81 s.” Across all electrodes, bins, bands "
     "and ROIs this yields about **25,000 candidate features per trial.** (Section 6 describes how the "
     "model trims this down to the informative handful.)")
figure(doc, "fig_binning.png", "How a continuous waveform becomes features. The grey line is the per-bin "
       "CNV signal at the vertex; the coloured bars are the per-bin summaries — each bar is one feature. "
       "The model uses 62.5 ms bins (finer than the 0.25 s bins drawn here for clarity).", 4)
para(doc, "Below are two examples built directly from the project’s own data, averaged across all 30 "
     "participants.")
figure(doc, "fig_cnv_waveform.png", "Grand-average CSD-transformed potential at the vertex (Cz), "
       "averaged over participants. The slow negative-going deflection across the foreperiod is the CNV. "
       "The straight (blue) and diagonal (red) conditions overlap almost completely: the "
       "class-discriminative variance is small relative to within- and between-subject variance, which "
       "is what makes single-trial decoding hard.", 5)
figure(doc, "fig_psd_bands.png", "Single-trial band-limited spectral power (Morlet) at Cz, "
       "grand-averaged. The 1/f-dominated profile (delta ≫ gamma) is typical of EEG. The two conditions "
       "are near-identical at the univariate level — consistent with discriminative information residing "
       "in distributed, higher-order feature combinations rather than any single band-power contrast.", 6)
callout(doc, "Why this matters",
        ["Notice that in Figures 5 and 6 the two conditions look nearly identical. That is the central "
         "challenge: there is no single univariate “tell” — no one electrode, latency, or band separates "
         "straight from diagonal. The model’s job is to find a faint, distributed pattern spread across "
         "thousands of features — precisely the needle-in-a-haystack task that XGBoost, combined with "
         "the feature-selection steps in Section 6, is designed for."])

# ----- 5 POOLING -----
heading(doc, "5   Borrowing strength across people: partial pooling", 1)
para(doc, "Every person’s brain is a little different, so the safest approach is to train a separate "
     "model for each participant using only their own trials. But there is a catch: with only ~80 trials "
     "per person, a flexible model can easily **“memorise” the training trials** instead of learning a "
     "general pattern — a problem called **overfitting**. It then looks great on the data it has seen and "
     "falls apart on new data.")
para(doc, "**Partial pooling** is the fix. The idea is to let each person’s model also learn from "
     "*everyone else’s* trials, while still being tested only on that person. There are three options on "
     "a spectrum:")
figure(doc, "fig_pooling.png", "Three ways to use the group’s data when predicting one person (marked "
       "with a star). Per-participant uses only that person’s data (prone to overfitting). Full pooling "
       "uses everyone except them. Partial pooling — the approach we use — trains on the target person "
       "plus everyone else, then tests on the target.", 7)
para(doc, "**What is combined, and how:** each participant’s trials are rows in a big table (with the "
     "~25,000 feature columns described above). Partial pooling simply **stacks everyone’s rows into one "
     "large training table** — about **2,400 trials** instead of ~80 — so the model sees the broad, "
     "shared pattern of how preparation differs between straight and diagonal steps. The held-out trials "
     "being scored still belong only to the one person we are predicting, so the test stays fair.")
callout(doc, "What partial pooling bought us",
        ["Pooling’s biggest benefit is **honesty**. Before pooling, the models showed a large gap between "
         "their optimistic internal scores and their true performance on unseen data — the signature of "
         "overfitting. After pooling, that gap essentially **disappears** (it even reverses slightly: see "
         "Section 7), while the headline accuracy holds up or modestly improves. In short, pooling makes "
         "the reported numbers trustworthy."], accent=GREEN, fill="F1F8F3")

# ----- 6 ARCHITECTURE -----
heading(doc, "6   How the model is built and tested, step by step", 1)
para(doc, "Two ideas run through the whole pipeline: **(a)** never let the model peek at the data it "
     "will be judged on, and **(b)** ruthlessly trim ~25,000 raw features down to the few dozen that "
     "genuinely matter.")
heading(doc, "An honest test: nested cross-validation", 2)
para(doc, "To measure performance fairly, we use **cross-validation**: the trials are split into five "
     "parts, and each part takes a turn as the “unseen test” while the model learns from the other four. "
     "We then do this in a **nested** way — an outer split used *only* for the final scoring, and an "
     "inner split used for all the tuning decisions. This separation is what stops the model from "
     "flattering itself.")
figure(doc, "fig_nestedcv.png", "Nested cross-validation. The outer split sets aside a test block used "
       "only to score the finished model. All the choices — which features to keep, which settings to "
       "use — are made on a separate inner split of the remaining data. Because the test block is never "
       "used to make any choice, the final score is not inflated.", 8)
heading(doc, "The feature funnel", 2)
para(doc, "Inside *each* training fold (so that no test data leaks in), the ~25,000 features pass through "
     "a sequence of filters, each removing features that are redundant or unhelpful, like a funnel "
     "narrowing to the essentials:")
figure(doc, "fig_pipeline.png", "The feature funnel, applied fresh inside every training fold. It starts "
       "from ~25,000 candidate features and ends with roughly 120 that survive every filter.", 9)
simple_table(doc, ["Stage", "What it does, in plain terms"], [
    ["**Correlation filter**", "Drops features that say almost the same thing as another feature, to avoid wasteful duplication."],
    ["**ANOVA screen**", "A quick statistical test keeps only the ~500 features most related to step type and discards the obvious non-starters."],
    ["**Stability selection**", "Re-runs the selection on many random sub-samples and keeps only features that repeatedly prove useful — not ones that looked good by luck once."],
    ["**Gain pruning**", "Trains an XGBoost model and drops any feature the trees never actually used to make a split."],
    ["**SHAP pruning**", "Measures how much each remaining feature really influenced the predictions and removes the least-influential 20%."],
], [2200, 7160])
doc.add_paragraph().paragraph_format.space_after = Pt(2)
heading(doc, "What is “SHAP”?", 2)
para(doc, "**SHAP** is a principled way to answer “*how much did each feature contribute to this "
     "prediction?*” It comes from game theory: it treats the features like players on a team and fairly "
     "divides the credit for the final prediction among them. Features that consistently earn little "
     "credit are pruned away. The pay-off is twofold — a leaner, less overfit model, and a ranked list "
     "of which brain measurements actually drive the decision, which is scientifically interesting in "
     "its own right.")
heading(doc, "Tuning the XGBoost settings", 2)
para(doc, "Finally, XGBoost has dials — how deep each tree can grow, how fast it learns, how strongly it "
     "is penalised for complexity. We search across many combinations (efficiently discarding the weak "
     "ones early) and keep the best, judged only on the inner tuning data. After the funnel, each final "
     "per-person model rests on roughly **120 features**.")

# ----- 7 RESULTS -----
heading(doc, "7   Current results", 1)
para(doc, "The numbers below come from the most recent **full-cohort run**: all 30 participants, partial "
     "pooling, the complete ~25,000-feature set, and the full funnel above. Performance is reported "
     "mainly as **AUC**, with accuracy alongside.")
callout(doc, "AUC vs. accuracy — what’s the difference?",
        ["**Accuracy** is the simplest score: the percentage of trials put in the right box, using one "
         "fixed cut-off (e.g. “call it diagonal if the model is more than 50% sure”). It is easy to read "
         "but can mislead — if one step type were more common, a lazy model could score well just by "
         "always guessing that one.",
         "**AUC** (the Area Under the ROC Curve) is a fuller summary. Rather than fixing a cut-off, it "
         "asks: *if you pick one real straight trial and one real diagonal trial at random, how often "
         "does the model give the diagonal one the higher score?* 1.00 means always, 0.50 means it is "
         "guessing. Because it sweeps across every possible cut-off, AUC is not fooled by an uneven mix "
         "of classes and rewards the model for *ranking* trials correctly even when it is unsure — which "
         "is why it is our headline number."])
metric_cards(doc, [("0.714", "average AUC across the 30 participants"),
                   ("67%", "average accuracy (straight vs diagonal)"),
                   ("0.47–1.00", "range across individuals")])
para(doc, "A **confusion matrix** is the simplest way to see *where* that accuracy comes from. It "
     "cross-tabulates what actually happened against what the model predicted: the green diagonal counts "
     "the correct calls, and the red off-diagonal counts the mistakes.")
figure(doc, "fig_confusion.png", "Confusion matrix for the whole cohort, pooling all 2,384 predictions. "
       "Rows are what the step actually was; columns are what the model predicted. Green cells (the "
       "diagonal) are correct calls; red cells are mistakes. The two step types are read about equally "
       "well — 66% of straight and 68% of diagonal steps correct — so the model is genuinely telling "
       "them apart rather than leaning toward one answer.", 10)
para(doc, "The errors are **balanced**: the model is roughly as good at spotting straight steps (66% "
     "correct) as diagonal ones (68% correct), with no strong bias toward either. The next chart breaks "
     "the same results down by participant.")
figure(doc, "fig_auc_bars.png", "Each participant’s own model, scored on their held-out trials and "
       "sorted from weakest to strongest. The dashed line is chance (0.50); the solid line is the group "
       "average (0.714). Most people (green) are read well above chance; a handful sit near chance, and "
       "one falls just below — a reminder that this works far better for some brains than others.", 11)
para(doc, "Two things stand out. First, **the model works, on average, clearly above chance**: 19 of 30 "
     "participants reach an AUC of 0.70 or higher, and several are read almost perfectly. Second, there "
     "is **wide individual variation** — a few participants sit near the chance line. Understanding why "
     "some brains are so much more “readable” than others is an open and interesting question.")
callout(doc, "The honesty check",
        ["A model that is overfitting scores higher on its own tuning data than on truly unseen data. "
         "Here the opposite holds: the average score on the **unseen** test trials (0.714) is actually "
         "*higher* than the internal tuning score (0.654). The gap between “optimistic internal” and "
         "“real-world” performance has not just closed — it has slightly reversed. This is the strongest "
         "sign that **the 0.71 figure is real and not inflated**, and it is the direct pay-off of partial "
         "pooling."], accent=GREEN, fill="F1F8F3")
callout(doc, "How to read this honestly",
        ["The group average (0.71) is the stable, trustworthy headline. **Individual** scores rest on "
         "small test sets (~16 trials per fold), so a single person’s value — including the perfect "
         "1.00 — is noisier and should be read as “high” rather than as an exact figure. Separately, "
         "earlier controlled comparisons that deliberately used a *trimmed* feature set (for speed) "
         "landed around 0.60–0.64; those experiments were what confirmed the overfitting-gap collapse "
         "described above. The full-feature run reported here is the project’s current best honest "
         "estimate."], accent=WARN, fill="FDF6EE")

# ----- 8 FEATURE INFORMATIVENESS -----
heading(doc, "8   Which features carry the signal?", 1)
para(doc, "Because the selection step (Section 6) runs inside all 150 training folds, we can simply tally "
     "**how often each feature is retained**. Features kept in every fold are the most robustly "
     "informative, and the pattern across them is physiologically revealing — and reassuringly "
     "consistent with the CNV literature.")
figure(doc, "fig_feat_informativeness.png", "Feature informativeness, measured as how often each feature "
       "is retained across the 150 training folds. (a) By type, the discriminative signal is carried "
       "overwhelmingly by temporal-slope features, with a smaller theta-power contribution; the binned "
       "amplitude and eLORETA source features, though available, are essentially never retained. (b) By "
       "timing, informativeness is concentrated in the early foreperiod, peaking ~0.2 s after the "
       "direction cue.", 12)
para(doc, "Two patterns stand out. **By feature type**, the discriminative signal is carried almost "
     "entirely by the **temporal-slope** features — the instantaneous rate of change of the potential — "
     "with a secondary contribution from **spectral power**, predominantly in the **theta** band "
     "(4–8 Hz). The per-bin amplitude and the source-space features were rarely selected. **By timing**, "
     "informativeness is sharply concentrated in the **early foreperiod (~0.1–0.35 s post-cue)**, peaking "
     "near 0.2 s, rather than in the terminal CNV. **Spatially**, the retained features concentrate over "
     "**central and fronto-central sensorimotor cortex** (the C, FC and CP electrode rows — Cz, C1–C6, "
     "FC1–FC4, CP4; strongest on the right, TP8/C6/C4). Taken together, the decoder is keying on an "
     "**early, sensorimotor, rate-of-change signal with a fronto-central theta component** — a "
     "physiologically sensible correlate of motor-plan formation, and a result of interest in its own "
     "right.")
para(doc, "To make these features concrete, the matrix below shows **example single-trial topographies** "
     "— different feature types (amplitude, frequency-band power, and eLORETA source), different "
     "participants, conditions and time bins. They are deliberately raw single trials: individually "
     "noisy, which is exactly why the selection funnel and cross-subject pooling matter.")
figure(doc, "fig_feature_topomaps.png", "Example single-trial feature topographies. The top-down scalp "
       "maps show amplitude (CSD) and band-power features at one electrode set; the two brain views are "
       "top-down glass-brains of the eLORETA source features (one dot per cortical region). Each panel is "
       "one real trial, labelled underneath with its feature, participant, condition (straight vs "
       "diagonal) and time bin, and scaled to its own range (red = higher, blue = lower). Single trials "
       "are noisy by nature — the patterns become reliable only after the selection and pooling described "
       "above.", 13, width_in=6.2)

# ----- 9 NEXT -----
heading(doc, "9   Summary and next steps", 1)
para(doc, "In plain terms: **the brain’s preparation signal does carry a readable trace of which way a "
     "person is about to step**, and a carefully-built, honestly-tested XGBoost model can pick it up — "
     "on average around 0.71 (with 0.50 being chance), and much higher in the most readable individuals. "
     "Partial pooling across participants was the key step that made these numbers trustworthy rather "
     "than optimistic. Encouragingly, the model also points to a clear, interpretable substrate — an "
     "early, sensorimotor, rate-of-change signal — rather than an opaque pattern.")
callout(doc, "Worth publishing",
        ["We think there is a **paper** in this. An honest, fully cross-validated decoding of step "
         "direction from the CNV — together with the finding that the signal is carried by early "
         "sensorimotor slope features — would make a nice contribution to the literature."],
        accent=GREEN, fill="F1F8F3")
para(doc, "**Where this can go next:**")
bullet(doc, "**Understanding individual differences** — why some participants are read almost perfectly and others barely above chance.")
bullet(doc, "**Localising the early signal** — pinning down the cortical source of the informative early-foreperiod slope features more precisely (the current template-based source features added little and could be improved).")
bullet(doc, "**Alternative models** — among the models we have evaluated, **XGBoost has been the most promising**: it gave the strongest results on small subsets of data and with much shorter training times, and it pairs that with a sophisticated, carefully-staged model architecture. The other candidates are modern neural-network models (compact “CNN/EEGNet” deep-learning networks designed for EEG); these are worth exploring as a next step, but may not ultimately improve on the current model.")

# ----- GLOSSARY -----
heading(doc, "Plain-language glossary", 1)
gloss = [
    ("EEG", "Electroencephalography — recording the brain’s electrical activity with sensors on the scalp."),
    ("CNV", "Contingent Negative Variation — a slow brain wave that builds up while a person prepares for an expected action."),
    ("Epoch / trial", "One ~2-second recording from a single attempt, across all sensors."),
    ("Feature", "A single number summarising one aspect of a trial (e.g. the average signal at one sensor in one time slice)."),
    ("Classification", "Predicting which category something belongs to — here, “straight” vs “diagonal.”"),
    ("AUC", "A performance score from 0.50 (random guessing) to 1.00 (perfect). 0.70 is a solidly useful model on hard biological data."),
    ("Overfitting", "When a model memorises its training data instead of learning a general rule, so it looks good in practice but fails on new data."),
    ("Cross-validation", "Repeatedly holding out part of the data as an unseen test to estimate real-world performance fairly."),
    ("Partial pooling", "Training each person’s model on their own data plus everyone else’s, to borrow the stability of the group while staying personalised."),
    ("XGBoost", "The model used here: hundreds of small decision trees added together, each correcting the previous ones’ mistakes."),
    ("SHAP", "A fair-credit method for measuring how much each feature contributed to a prediction; used to prune weak features and to interpret the model."),
]
for term, defn in gloss:
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    rt = p.add_run(term + " — "); rt.bold = True; rt.font.color.rgb = RGBColor.from_string('16324F')
    rd = p.add_run(defn)

set_update_fields(doc)
doc.save(str(OUT))
print("DOCX written:", OUT, "(%.0f KB)" % (OUT.stat().st_size / 1024))
